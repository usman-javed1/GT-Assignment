from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_title_page(doc, student_name="Usman Javed", reg_number="2022CS620"):
    """Create the title page of the report"""
    # Add title
    title = doc.add_paragraph()
    title_run = title.add_run("Page Blocks Classification Dataset Analysis")
    title_run.font.size = Pt(24)
    title_run.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add student info
    doc.add_paragraph()  # Add some space
    student_info = doc.add_paragraph()
    student_info.add_run(f"{student_name}\n{reg_number}")
    student_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add course info
    doc.add_paragraph()  # Add some space
    course_info = doc.add_paragraph()
    course_info.add_run("Data Science Course Project\n6th Semester")
    course_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add page break
    doc.add_page_break()

def add_section(doc, title, content):
    """Add a section with title and content"""
    heading = doc.add_heading(title, level=1)
    doc.add_paragraph(content)

def add_figure(doc, image_path, caption):
    """Add a figure with caption"""
    doc.add_picture(image_path, width=Inches(6))
    caption_paragraph = doc.add_paragraph(caption)
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()  # Add some space

def generate_report():
    """Generate the complete report"""
    doc = Document()
    
    # Create title page
    create_title_page(doc)
    
    # Introduction
    add_section(doc, "Introduction", 
                "This report presents an analysis of the Page Blocks Classification dataset from the UCI Machine Learning Repository. "
                "The dataset consists of features extracted from document page layouts, with the goal of classifying different types "
                "of blocks (text, horizontal line, picture, vertical line, and graphic) in the page layout.")
    
    # Data Description
    add_section(doc, "Dataset Description",
                "The dataset contains 5473 examples with 10 features and 1 target variable. The features include physical attributes "
                "of the blocks such as height, length, area, and various pixel-based measurements. The target variable represents "
                "five different classes of page layout elements.")
    
    # Univariate Analysis
    add_section(doc, "Univariate Analysis", 
                "The univariate analysis examines individual features independently to understand their distributions and characteristics.")
    
    # Add univariate plots
    add_figure(doc, "../images/univariate_height.png", 
               "Figure 1: Distribution of block heights showing a right-skewed distribution with most blocks having small heights.")
    add_figure(doc, "../images/univariate_class.png",
               "Figure 2: Distribution of block classes showing class imbalance with text blocks being the most common.")
    
    # Bivariate Analysis
    add_section(doc, "Bivariate Analysis",
                "The bivariate analysis explores relationships between pairs of features and their associations with the target class.")
    
    # Add bivariate plots
    add_figure(doc, "../images/bivariate_height_length.png",
               "Figure 3: Relationship between block height and length, colored by class, showing distinct clustering patterns.")
    add_figure(doc, "../images/bivariate_black_area.png",
               "Figure 4: Relationship between block area and black pixels, revealing strong correlation and class-specific patterns.")
    
    # Multivariate Analysis
    add_section(doc, "Multivariate Analysis",
                "The multivariate analysis examines relationships between multiple features simultaneously to uncover complex patterns.")
    
    # Add multivariate plots
    add_figure(doc, "../images/multivariate_pairplot.png",
               "Figure 5: Pairwise relationships between key features showing complex interactions and class separation.")
    add_figure(doc, "../images/multivariate_correlation.png",
               "Figure 6: Correlation heatmap of numeric features highlighting strong relationships between certain features.")
    
    # Conclusions
    add_section(doc, "Conclusions",
                "The analysis reveals several important insights about the page blocks dataset:\n"
                "1. There is significant class imbalance with text blocks being the most common type.\n"
                "2. Block dimensions (height, length, area) show distinct patterns for different classes.\n"
                "3. Strong correlations exist between physical measurements and pixel-based features.\n"
                "4. The features provide good separation between different block classes, suggesting potential for accurate classification.")
    
    # Save the document
    doc.save("../output/2022CS620.docx")

if __name__ == "__main__":
    generate_report() 